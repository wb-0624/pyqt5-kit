import docx


class WordKit:
    def __init__(self, path):
        self.word_path = path
        self.word_name = path.split('/')[-1]
        self.word = None

    def read_word(self):
        if self.word_name.endswith('.doc'):
            self.read_doc()
        elif self.word_name.endswith('.docx'):
            self.read_docx()
        else:
            raise Exception('File type error!')

    def read_doc(self):
        raise Exception('doc is not supported!')

    def read_docx(self):
        self.word = docx.Document(self.word_path)
        text_list = []
        for p in self.word.paragraphs:
            text_list.append(p.text)
        result = '\n'.join(text_list)
        return result

    def write_word(self, data):
        if self.word_name.endswith('.doc'):
            self.write_doc(data)
        elif self.word_name.endswith('.docx'):
            self.write_docx(data)
        else:
            raise Exception('File type error!')

    def write_doc(self, data):
        raise Exception('doc is not supported!')

    def write_docx(self, data: list):
        if self.word is None:
            self.word = docx.Document()
        for p in data:
            self.word.add_paragraph(p)

    def save(self):
        self.word.save(self.word_path)


if __name__ == '__main__':
    word = WordKit('2020.docx')
    word.write_word(['1', '2', '3'])
    word.save()
